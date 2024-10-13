import json

from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image

import subprocess
import os


class LaptopDocsDataSource:

    def createDocx(self, data, id):
        doc = Document()
        doc.add_heading(f"Серийный номер: {data['title']}", level=1)
        doc.add_heading(f"Протокол #{id}", level=2)

        # Добавляем описание (description)
        doc.add_paragraph('12.02.2005')

        doc.add_heading(f"Описание", level=2)
        doc.add_paragraph(f'Дефекты: {len(data["deffects"])-1}')
        for i in range(len(data["deffects"])-1):
            doc.add_heading(f"Дефект{i+1}: Описание", level=3)
            item = json.loads(data["deffects"][i])
            doc.add_paragraph(item["description"])
            image_path = f'out_image/{item["id"]}.jpg'  # путь к вашему изображению

            doc.add_picture(image_path,width=Inches(3), height=Inches(2))

        # Добавляем изображение

        doc.save(f'docx_output/{id}.docx')
    def createPdf(self, data, id):
        self.createDocx(data, id)
        output_dir = f'pdf_output/'
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
    
        command = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, f'docx_output/{id}.docx']
        try:
            subprocess.run(command, check=True)
            print(f"Conversion successful! PDF saved to {output_dir}")
        except subprocess.CalledProcessError as e:
            print(f"Error during conversion: {e}")
